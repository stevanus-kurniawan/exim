"""
Extract text from .docx files in Docs folder and save as .md for review.
"""
import os
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCS_SRC = ROOT / "Docs"
OUT_DIR = ROOT / "docs"

def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
        parts.append("")  # newline after table
    return "\n\n".join(parts)

def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if not DOCS_SRC.exists():
        print(f"Docs folder not found: {DOCS_SRC}")
        return
    for f in DOCS_SRC.glob("*.docx"):
        if f.name.startswith("~$"):
            continue
        text = extract_docx(f)
        out_name = f.stem + ".md"
        out_path = OUT_DIR / out_name
        with open(out_path, "w", encoding="utf-8") as out:
            out.write(f"# {f.stem}\n\n")
            out.write(text)
        print(f"Extracted: {f.name} -> {out_path}")

if __name__ == "__main__":
    main()
